"""
Docx生成サービス

Markdownから編集可能なDocxファイルを生成
"""

import logging
import re
from typing import Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class DocxGenerator:
    """Docxファイル生成"""

    def generate_docx_from_markdown(
        self,
        markdown_text: str,
        layout_metadata: Optional[dict] = None,
        target_language: str = 'ja',
        job_id: Optional[str] = None
    ) -> bytes:
        """
        Markdownから直接Docxを生成

        Args:
            markdown_text: Markdownテキスト
            layout_metadata: レイアウト情報（未使用）
            target_language: 言語コード
            job_id: ジョブID

        Returns:
            Docxのバイト列
        """
        from app.services.document_preprocessor import DocumentPreprocessor
        import io

        # Docx用にMarkdownを前処理（ページ番号削除 + 改ページマーカー挿入）
        preprocessor = DocumentPreprocessor()
        docx_markdown = preprocessor.prepare_for_paged_output(
            markdown_text,
            output_format='docx'
        )

        # Docx文書を作成
        doc = Document()

        # デフォルトフォントを設定
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # 画像ディレクトリのパスを取得
        figures_dir = None
        if job_id:
            figures_dir = Path("storage").resolve() / "documents" / job_id / "figures"
            logger.info(f"Docx generation: Looking for images in {figures_dir}")

        # Markdownを行ごとに処理
        lines = docx_markdown.split('\n')
        i = 0
        image_count = 0
        found_count = 0
        missing_count = 0
        in_list = False
        list_items = []

        while i < len(lines):
            line = lines[i]

            # 改ページマーカーを検出
            if '<!-- PAGE_BREAK -->' in line:
                doc.add_page_break()
                i += 1
                continue

            # 見出し処理（フォントサイズを調整）
            if line.startswith('# '):
                text = line[2:].strip()
                # H1: 14pt（問題番号など）
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(14)
                run.font.bold = True
                i += 1
                continue
            elif line.startswith('## '):
                text = line[3:].strip()
                # H2: 12pt
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(12)
                run.font.bold = True
                i += 1
                continue
            elif line.startswith('### '):
                text = line[4:].strip()
                # H3: 11pt + Bold（設問部分）
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(11)
                run.font.bold = True
                i += 1
                continue
            elif line.startswith('#### '):
                text = line[5:].strip()
                # H4: 11pt + Bold
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(11)
                run.font.bold = True
                i += 1
                continue

            # 画像処理
            img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)
                image_count += 1

                # 相対パスから実際のファイルパスを取得
                if img_path.startswith('figures/') and figures_dir:
                    filename = img_path.replace('figures/', '')
                    full_path = figures_dir / filename

                    if full_path.exists():
                        try:
                            # 画像を挿入（幅を6インチに制限）
                            doc.add_picture(str(full_path), width=Inches(6))
                            found_count += 1
                            logger.info(f"✓ Image inserted: {filename}")
                        except Exception as e:
                            logger.error(f"Failed to insert image {filename}: {e}")
                            # 画像挿入失敗時はキャプションのみ追加
                            p = doc.add_paragraph(f"[Image: {alt_text}]")
                            p.italic = True
                            missing_count += 1
                    else:
                        logger.warning(f"✗ Image NOT found: {filename} (expected at {full_path})")
                        # 画像が見つからない場合はキャプションのみ追加
                        p = doc.add_paragraph(f"[Image: {alt_text}]")
                        p.italic = True
                        missing_count += 1
                else:
                    # 画像パスが不正な場合
                    p = doc.add_paragraph(f"[Image: {alt_text}]")
                    p.italic = True
                    missing_count += 1

                i += 1
                continue

            # リスト項目の検出（順序なしリスト: - または * で始まる）
            list_match = re.match(r'^[-*]\s+(.+)$', line.strip())
            if list_match:
                if not in_list:
                    in_list = True
                    list_items = []
                list_items.append(list_match.group(1))
                i += 1
                continue
            elif in_list and not line.strip():
                # リストが終了（空行）
                if list_items:
                    self._add_two_column_list(doc, list_items)
                    list_items = []
                in_list = False
                doc.add_paragraph()
                i += 1
                continue
            elif in_list:
                # リスト終了（リスト項目ではない行）
                if list_items:
                    self._add_two_column_list(doc, list_items)
                    list_items = []
                in_list = False
                # この行は次のループで処理
                continue

            # 通常のテキスト処理
            if line.strip():
                # 太字、イタリックなどの簡易処理
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
            else:
                # 空行
                doc.add_paragraph()

            i += 1

        # ループ終了時にリストが残っている場合
        if in_list and list_items:
            self._add_two_column_list(doc, list_items)

        logger.info(f"Docx image processing complete: {image_count} total, {found_count} found, {missing_count} missing")

        # メモリ上にDocxを保存
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        return docx_bytes.read()

    def _add_two_column_list(self, doc, items: list):
        """
        リスト項目を適切な列数の表として追加
        選択肢の数とテキスト長を考慮して列数を決定

        Args:
            doc: python-docxのDocumentオブジェクト
            items: リスト項目のリスト
        """
        # 列数を決定
        num_columns = self._determine_list_columns(items)

        # 項目数に応じて行数を計算
        rows = (len(items) + num_columns - 1) // num_columns

        # 表を作成
        table = doc.add_table(rows=rows, cols=num_columns)
        table.style = 'Table Grid'

        # 各セルに項目を配置
        for idx, item in enumerate(items):
            row_idx = idx // num_columns
            col_idx = idx % num_columns
            cell = table.cell(row_idx, col_idx)

            # セルのテキストを設定
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(item)
            run.font.size = Pt(11)

            # セルの余白を設定
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')

            for margin_name in ['top', 'left', 'bottom', 'right']:
                node = OxmlElement(f'w:{margin_name}')
                node.set(qn('w:w'), '100')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)

            tcPr.append(tcMar)

        # 表の後に空行を追加
        doc.add_paragraph()

        logger.info(f"Added {num_columns}-column list with {len(items)} items ({rows} rows)")

    def _determine_list_columns(self, items: list) -> int:
        """
        リスト項目の数とテキスト長から最適な列数を決定

        Args:
            items: リスト項目のリスト

        Returns:
            列数（1, 2, または 3）
        """
        num_items = len(items)

        # 選択肢が2個以下なら1列
        if num_items <= 2:
            logger.info(f"List layout: 1 column (only {num_items} items)")
            return 1

        # 平均文字数を計算
        avg_length = sum(len(item) for item in items) / len(items)
        max_length = max(len(item) for item in items)

        logger.info(f"List analysis: {num_items} items, avg={avg_length:.1f} chars, max={max_length} chars")

        # テキスト長に基づいて列数を決定
        if max_length > 80 or avg_length > 60:
            # 非常に長い選択肢 → 1列
            logger.info("List layout: 1 column (text too long)")
            return 1
        elif max_length > 50 or avg_length > 40:
            # 中程度の長さ → 2列まで
            if num_items <= 4:
                logger.info("List layout: 2 columns (medium length, few items)")
                return 2
            else:
                logger.info("List layout: 1 column (medium length, many items)")
                return 1
        else:
            # 短い選択肢 → 積極的に2-3列
            if num_items <= 6:
                logger.info("List layout: 2 columns (short text)")
                return 2
            else:
                # 7個以上の短い選択肢 → 3列も検討
                logger.info("List layout: 3 columns (short text, many items)")
                return 3

    def _add_formatted_text(self, paragraph, text: str):
        """
        簡易的なMarkdown書式をDocxに変換

        Args:
            paragraph: python-docxの段落オブジェクト
            text: Markdownテキスト
        """
        # 簡易実装：太字（**text**）とイタリック（*text*）のみ対応
        # より高度な処理は必要に応じて拡張

        # 現時点ではプレーンテキストとして追加
        paragraph.add_run(text)
